from __future__ import annotations

import argparse
import csv
import json
import os
import shutil
import subprocess
import sys
import tempfile
import time
from contextlib import ExitStack
from pathlib import Path
from statistics import median
from unittest.mock import patch

REPO_ROOT = Path(__file__).resolve().parents[2]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))


def _legacy_parser_worker(path: str) -> dict:
    from backend.app.core import quarantine

    command = [sys.executable, "-m", "backend.app.core.parser_worker", path]
    try:
        completed = subprocess.run(
            command,
            check=False,
            capture_output=True,
            timeout=quarantine.PARSER_TIMEOUT_SECONDS,
            env=quarantine._worker_env(),
        )
    except subprocess.TimeoutExpired as exc:
        raise quarantine.QuarantineError("Parser worker timed out") from exc
    stdout_text = quarantine._decode_worker_stream(completed.stdout)
    stderr_text = quarantine._decode_worker_stream(completed.stderr)
    if completed.returncode != 0:
        detail = (stderr_text or stdout_text or "Parser worker failed").strip()
        raise quarantine.QuarantineError(detail[:500])
    raw = stdout_text.encode("utf-8")
    if len(raw) > quarantine.MAX_WORKER_JSON_BYTES:
        raise quarantine.QuarantineError("Parser worker output exceeded the allowed size")
    try:
        payload = json.loads(stdout_text)
    except json.JSONDecodeError as exc:
        raise quarantine.QuarantineError("Parser worker returned malformed JSON") from exc
    return quarantine.validate_worker_output(payload)


def _legacy_ingest(vault_id: str, path: str) -> dict:
    """Pre-hardening equivalent: same AV scan, without new fail-closed/containment controls."""
    from backend.app.core import quarantine

    candidate = quarantine.validate_candidate_file(path)
    defender = quarantine.defender_scan(candidate["canonical_path"])
    encrypted_blob = quarantine._encrypted_quarantine_blob(
        vault_id, Path(candidate["canonical_path"])
    )
    record_id = quarantine.create_quarantine_record(
        vault_id, candidate, defender, encrypted_blob=encrypted_blob
    )
    try:
        parsed = quarantine.parse_candidate_file(candidate)
        quarantine.update_quarantine_record(
            record_id,
            validation_status="passed",
            parser_status="passed",
            parser_detail="worker_output_validated",
            trust_tier=quarantine._trust_tier(candidate, defender),
        )
    except Exception as exc:
        quarantine.update_quarantine_record(
            record_id,
            validation_status="passed",
            parser_status="failed",
            parser_detail=str(exc)[:500],
            trust_tier="quarantined",
        )
        raise
    parsed["quarantine_record_id"] = record_id
    parsed["security"] = {
        "validation": candidate,
        "defender": defender,
        "encrypted_blob": encrypted_blob,
        "trust_tier": quarantine._trust_tier(candidate, defender),
        "provenance": "local_import",
        "security_labels": ["local", "quarantined-parse"],
        "parser": parsed.get("parser") or {},
    }
    return parsed


def _create_corpus(root: Path, count: int) -> list[str]:
    from docx import Document

    root.mkdir(parents=True, exist_ok=True)
    template = root / "template.docx"
    document = Document()
    document.add_heading("Vault ingestion security benchmark", level=1)
    document.add_paragraph(
        "A small representative document used to measure parser isolation, "
        "quarantine validation, durable storage, and batch ingestion throughput."
    )
    document.save(template)
    paths: list[str] = []
    for index in range(count):
        target = root / f"document-{index:05d}.docx"
        shutil.copyfile(template, target)
        paths.append(str(target))
    template.unlink()
    return paths


def _run_once(corpus: list[str], mode: str, run_root: Path, *, worker_count: int) -> dict:
    os.environ["CML_DATA_DIR"] = str(run_root / "data")
    os.environ["CML_DATABASE_PATH"] = str(run_root / "vault.sqlite3")
    os.environ["CML_EMBEDDING_PROVIDER"] = "hash"
    os.environ["CML_ALLOW_HASH_EMBEDDINGS"] = "1"

    from backend.app.core.config import get_settings

    get_settings.cache_clear()
    from backend.app.api.routes.sources import create_source_import_job
    from backend.app.core.background_jobs import run_due_jobs_once
    from backend.app.core.database import connect, init_db, utc_now
    from backend.app.core.migrations import run_migrations
    from backend.app.schemas import SourceImportJobRequest

    init_db()
    run_migrations()
    now = utc_now()
    with connect() as conn:
        conn.execute(
            "INSERT INTO vaults (id, name, path, created_at, updated_at) VALUES (?, ?, ?, ?, ?)",
            ("benchmark-vault", "Benchmark", str(run_root), now, now),
        )

    # Make filesystem cache state comparable while leaving parsing and DB work real.
    for raw_path in corpus:
        Path(raw_path).read_bytes()

    clean_defender = {
        "status": "passed",
        "detail": "Controlled benchmark: identical clean Defender result in both arms.",
    }
    with ExitStack() as stack:
        stack.enter_context(
            patch(
                "backend.app.core.background_jobs.source_import_worker_count",
                side_effect=lambda total: max(1, min(total, worker_count)),
            )
        )
        stack.enter_context(
            patch("backend.app.core.quarantine.defender_scan", return_value=clean_defender)
        )
        if mode == "baseline":
            stack.enter_context(
                patch("backend.app.core.quarantine.run_parser_worker", _legacy_parser_worker)
            )
            stack.enter_context(
                patch(
                    "backend.app.api.routes.sources.ingest_file_through_quarantine",
                    _legacy_ingest,
                )
            )
            stack.enter_context(
                patch("backend.app.core.security_scans.enqueue_due_security_scan", return_value=None)
            )
        started = time.perf_counter()
        job = create_source_import_job(
            SourceImportJobRequest(vault_id="benchmark-vault", paths=corpus)
        )
        processed_jobs = run_due_jobs_once(limit=1)
        elapsed = time.perf_counter() - started

    with connect() as conn:
        stored = conn.execute(
            "SELECT status, result_json FROM app_jobs WHERE id = ?", (job["id"],)
        ).fetchone()
        source_count = int(
            conn.execute("SELECT COUNT(*) FROM sources").fetchone()[0]
        )
    progress = json.loads(stored["result_json"] or "{}")
    completed = int(progress.get("completed_files") or 0)
    failed = int(progress.get("failed_files") or 0)
    if processed_jobs != 1 or stored["status"] != "succeeded" or completed != len(corpus) or failed:
        raise RuntimeError(
            f"Invalid {mode} run: jobs={processed_jobs} status={stored['status']} "
            f"completed={completed} failed={failed} sources={source_count}"
        )
    return {
        "mode": mode,
        "file_count": len(corpus),
        "elapsed_seconds": elapsed,
        "files_per_second": len(corpus) / elapsed,
        "milliseconds_per_file": elapsed * 1000 / len(corpus),
        "worker_count": max(1, min(len(corpus), worker_count)),
        "completed_files": completed,
        "failed_files": failed,
        "source_count": source_count,
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--sizes", default="10,100,1000")
    parser.add_argument("--small-repetitions", type=int, default=3)
    parser.add_argument("--workers", type=int, default=4)
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("output/benchmarks/security-ingestion.json"),
    )
    args = parser.parse_args()
    sizes = [int(value) for value in args.sizes.split(",") if int(value) > 0]
    all_runs: list[dict] = []
    with tempfile.TemporaryDirectory(prefix="cml-security-ingestion-") as raw_root:
        root = Path(raw_root)
        corpus = _create_corpus(root / "corpus", max(sizes))
        for size in sizes:
            repetitions = args.small_repetitions if size < max(sizes) else 1
            for repetition in range(repetitions):
                order = ("baseline", "hardened") if repetition % 2 == 0 else ("hardened", "baseline")
                for mode in order:
                    run_root = root / f"run-{size}-{repetition}-{mode}"
                    run = _run_once(
                        corpus[:size],
                        mode,
                        run_root,
                        worker_count=max(1, args.workers),
                    )
                    run["repetition"] = repetition + 1
                    all_runs.append(run)
                    print(json.dumps(run), flush=True)

    summary: list[dict] = []
    for size in sizes:
        baseline = [row["elapsed_seconds"] for row in all_runs if row["file_count"] == size and row["mode"] == "baseline"]
        hardened = [row["elapsed_seconds"] for row in all_runs if row["file_count"] == size and row["mode"] == "hardened"]
        baseline_seconds = median(baseline)
        hardened_seconds = median(hardened)
        summary.append(
            {
                "file_count": size,
                "repetitions": len(baseline),
                "baseline_seconds": baseline_seconds,
                "hardened_seconds": hardened_seconds,
                "baseline_files_per_second": size / baseline_seconds,
                "hardened_files_per_second": size / hardened_seconds,
                "delta_seconds": hardened_seconds - baseline_seconds,
                "slowdown_percent": ((hardened_seconds / baseline_seconds) - 1) * 100,
            }
        )

    report = {
        "benchmark": "security-ingestion-overhead-v1",
        "corpus": "small DOCX files; parser isolation exercised for every file",
        "defender_control": "identical mocked clean pass; Defender scanning predates the compared fixes",
        "fixed_worker_count": max(1, args.workers),
        "sizes": sizes,
        "summary": summary,
        "runs": all_runs,
    }
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(report, indent=2), encoding="utf-8")
    csv_path = args.output.with_suffix(".csv")
    with csv_path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(summary[0]))
        writer.writeheader()
        writer.writerows(summary)
    print(json.dumps({"report": str(args.output), "csv": str(csv_path), "summary": summary}, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
