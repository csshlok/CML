from __future__ import annotations

import csv
import json
import textwrap
from pathlib import Path


DEFAULT_COUNTS = {
    "pdf": 6,
    "docx": 2,
    "md": 2,
    "txt": 1,
    "html": 2,
    "json": 1,
    "csv": 1,
}

SCALE_COUNTS = {
    "pdf": 120,
    "docx": 60,
    "md": 2200,
    "txt": 2200,
    "html": 2200,
    "json": 2200,
    "csv": 1020,
}

THEMES = [
    "parser reliability",
    "browser capture simplicity",
    "repeat-turn token savings",
    "OCR backlog recovery",
    "invoice table fidelity",
    "customer onboarding friction",
    "retrieval trust calibration",
    "knowledge base exports",
    "search analytics drift",
    "field operations notes",
]

TEAMS = [
    "research",
    "ops",
    "sales",
    "finance",
    "support",
    "growth",
    "platform",
    "security",
]

REGIONS = ["north", "south", "east", "west", "central"]


def create_synthetic_user_corpus(
    root: str | Path,
    *,
    counts: dict[str, int] | None = None,
    target_file_count: int | None = None,
) -> dict:
    base = Path(root)
    base.mkdir(parents=True, exist_ok=True)
    directories = {
        "pdf": base / "pdf",
        "docx": base / "docs",
        "md": base / "notes",
        "txt": base / "notes",
        "html": base / "exports",
        "json": base / "data",
        "csv": base / "data",
    }
    for target in set(directories.values()):
        target.mkdir(parents=True, exist_ok=True)

    requested = _normalize_counts(counts, target_file_count=target_file_count)
    all_paths: list[Path] = []
    pdf_paths: list[Path] = []
    queries = _benchmark_queries()

    for index in range(requested["pdf"]):
        path = directories["pdf"] / f"artifact-{index:05d}.pdf"
        pdf_paths.append(_create_pdf_variant(path, index))
    all_paths.extend(pdf_paths)

    for index in range(requested["docx"]):
        path = directories["docx"] / f"document-{index:05d}.docx"
        all_paths.append(_create_docx_variant(path, index))
    for index in range(requested["md"]):
        path = directories["md"] / f"note-{index:05d}.md"
        all_paths.append(_write_markdown(path, _markdown_body(index)))
    for index in range(requested["txt"]):
        path = directories["txt"] / f"log-{index:05d}.txt"
        all_paths.append(_write_text(path, _text_body(index)))
    for index in range(requested["html"]):
        path = directories["html"] / f"export-{index:05d}.html"
        all_paths.append(_write_html(path, _html_title(index), _html_body(index)))
    for index in range(requested["json"]):
        path = directories["json"] / f"snapshot-{index:05d}.json"
        all_paths.append(_write_json(path, _json_payload(index)))
    for index in range(requested["csv"]):
        path = directories["csv"] / f"metrics-{index:05d}.csv"
        all_paths.append(_write_csv(path, rows=_csv_rows(index)))

    return {
        "root": str(base),
        "counts": requested,
        "pdf_paths": [str(path) for path in pdf_paths],
        "all_paths": [str(path) for path in all_paths],
        "queries": queries,
        "expected_suffixes": sorted({path.suffix.lower() for path in all_paths}),
    }


def _normalize_counts(counts: dict[str, int] | None, *, target_file_count: int | None) -> dict[str, int]:
    requested = dict(DEFAULT_COUNTS if counts is None else counts)
    for key in ("pdf", "docx", "md", "txt", "html", "json", "csv"):
        requested[key] = max(0, int(requested.get(key, 0)))
    if target_file_count is None:
        return requested
    requested = _scale_counts_to_target(requested, target_file_count)
    return requested


def _scale_counts_to_target(seed_counts: dict[str, int], target_file_count: int) -> dict[str, int]:
    if target_file_count <= sum(seed_counts.values()):
        return dict(seed_counts)
    if target_file_count >= 10_000 and sum(seed_counts.values()) <= 20:
        counts = dict(SCALE_COUNTS)
        if sum(counts.values()) >= target_file_count:
            return counts
    counts = dict(seed_counts)
    text_order = ["csv", "json", "html", "txt", "md"]
    while sum(counts.values()) < target_file_count:
        for key in text_order:
            counts[key] += 1
            if sum(counts.values()) >= target_file_count:
                break
    return counts


def _benchmark_queries() -> list[str]:
    return [
        "Which files mention repeat-turn token savings or cheaper follow-up questions?",
        "Find evidence about PDF table fidelity, invoices, or parser layout quality.",
        "What says browser capture should stay simple, one-click, or desktop-provisioned?",
        "Which artifacts mention OCR backlog, screenshot indexing slowdown, or retry floods?",
        "Show the implementation target for repeat-turn token reduction and cache-warm savings.",
        "Which export mentions retrieval trust calibration or grounded citations?",
        "Find any field report about parser worker saturation and recovery throttling.",
        "What does the knowledge base say about rejecting empty retrieval-hit benchmarks?",
    ]


def _theme(index: int) -> str:
    return THEMES[index % len(THEMES)]


def _team(index: int) -> str:
    return TEAMS[index % len(TEAMS)]


def _region(index: int) -> str:
    return REGIONS[index % len(REGIONS)]


def _campaign(index: int) -> str:
    return f"CMP-{1000 + index}"


def _layout_name(index: int) -> str:
    names = [
        "two-column brief",
        "invoice grid",
        "research digest",
        "scan-like field report",
        "timeline itinerary",
        "knowledge base export",
        "meeting agenda",
        "risk register",
        "product memo",
        "faq sheet",
    ]
    return names[index % len(names)]


def _create_pdf_variant(path: Path, index: int) -> Path:
    family = index % 10
    if family == 0:
        return _pdf_two_column(path, index)
    if family == 1:
        return _pdf_invoice_grid(path, index)
    if family == 2:
        return _pdf_research_digest(path, index)
    if family == 3:
        return _pdf_scan_like(path, index)
    if family == 4:
        return _pdf_timeline(path, index)
    if family == 5:
        return _pdf_knowledge_export(path, index)
    if family == 6:
        return _pdf_meeting_agenda(path, index)
    if family == 7:
        return _pdf_risk_register(path, index)
    if family == 8:
        return _pdf_product_memo(path, index)
    return _pdf_faq_sheet(path, index)


def _pdf_two_column(path: Path, index: int) -> Path:
    doc = _new_pdf()
    page = doc.new_page(width=595, height=842)
    page.insert_text((44, 48), f"Strategy Brief {_campaign(index)}", fontsize=22)
    page.insert_text((44, 76), f"Theme: {_theme(index)} | Region: {_region(index)}", fontsize=11)
    left = textwrap.fill(
        f"The {_team(index)} team reported that {_theme(index)} is driving benchmark work. "
        "The first answer can be heavier if repeated questions become materially cheaper and stay grounded in citations. "
        "Desktop-provisioned browser capture remains easier to support than manual extension setup.",
        width=42,
    )
    right = textwrap.fill(
        f"This file uses the {_layout_name(index)} layout family. "
        "Parser selection matters because invoices, scans, notes, and exports all break in different ways. "
        "The benchmark target is honest measurement of first-turn quality and cache-warm token reduction.",
        width=42,
    )
    page.insert_textbox((44, 118, 272, 730), left, fontsize=11, lineheight=1.35)
    page.insert_textbox((316, 118, 546, 730), right, fontsize=11, lineheight=1.35)
    page.draw_line((296, 108), (296, 730), width=1)
    _save_pdf(doc, path)
    return path


def _pdf_invoice_grid(path: Path, index: int) -> Path:
    doc = _new_pdf()
    page = doc.new_page(width=595, height=842)
    page.insert_text((48, 48), f"Vendor Invoice {_campaign(index)}", fontsize=20)
    page.insert_text((48, 74), f"Team: {_team(index)}", fontsize=11)
    page.insert_text((360, 74), f"Theme: {_theme(index)}", fontsize=11)
    top = 120
    row_height = 34
    columns = [48, 255, 350, 445, 542]
    headers = ["Description", "Hours", "Rate", "Amount"]
    for x in columns:
        page.draw_line((x, top), (x, top + (row_height * 7)), width=1)
    for row_index in range(8):
        y = top + (row_index * row_height)
        page.draw_line((48, y), (542, y), width=1)
    for idx, header in enumerate(headers):
        page.insert_text((columns[idx] + 8, top + 22), header, fontsize=10)
    rows = [
        (f"{_theme(index)} review", "8", "145", "1160"),
        ("OCR backlog drill", "5", "145", "725"),
        ("Extension setup reduction", "6", "145", "870"),
        ("Warm-cache savings study", "7", "145", "1015"),
        ("Retrieval evidence QA", "4", "145", "580"),
        ("Invoice table fidelity test", "3", "145", "435"),
    ]
    for row_index, row in enumerate(rows, start=1):
        y = top + (row_index * row_height) + 20
        for col_index, value in enumerate(row):
            page.insert_text((columns[col_index] + 8, y), value, fontsize=10)
    page.insert_text(
        (48, 410),
        f"Invoice note {_campaign(index)}: preserve table layout during PDF extraction so line items remain retrievable.",
        fontsize=11,
    )
    _save_pdf(doc, path)
    return path


def _pdf_research_digest(path: Path, index: int) -> Path:
    doc = _new_pdf()
    page = doc.new_page(width=595, height=842)
    page.insert_text((46, 48), f"Retrieval Research Digest {_campaign(index)}", fontsize=21)
    sections = [
        ("Abstract", f"The {_theme(index)} benchmark compares first-turn packet reduction with warm-cache savings."),
        ("Finding", "Reference compression can cut tokens hard, but retrieval hits must stay inspectable and non-zero."),
        ("Implication", f"Query sets for {_team(index)} should include real retrieval-backed artifacts, not empty packet scaffolding."),
        ("Recommendation", "Reject any benchmark run with zero indexed chunks or zero retrieval hits before graphs are rendered."),
    ]
    y = 100
    for title, body in sections:
        page.insert_text((48, y), title, fontsize=14)
        page.insert_textbox((48, y + 12, 544, y + 88), textwrap.fill(body, width=80), fontsize=11, lineheight=1.3)
        y += 120
    _save_pdf(doc, path)
    return path


def _pdf_scan_like(path: Path, index: int) -> Path:
    source = _new_pdf()
    page = source.new_page(width=595, height=842)
    page.draw_rect((38, 38, 557, 804), width=1)
    page.insert_text((56, 74), f"Field Report {_campaign(index)}", fontsize=18)
    body = textwrap.fill(
        f"The {_theme(index)} queue backed up after repeated screenshot uploads in {_region(index)}. "
        "Operators saw stale progress indicators and retry floods. "
        "A throttled recovery policy would have reduced parser worker saturation and OCR backlog amplification.",
        width=56,
    )
    page.insert_textbox((56, 112, 520, 730), body, fontsize=15, lineheight=1.6)
    pix = page.get_pixmap(dpi=180, alpha=False)
    raster = _new_pdf()
    out = raster.new_page(width=595, height=842)
    out.insert_image(out.rect, stream=pix.tobytes("png"))
    _save_pdf(raster, path)
    source.close()
    return path


def _pdf_timeline(path: Path, index: int) -> Path:
    doc = _new_pdf()
    page = doc.new_page(width=595, height=842)
    page.insert_text((48, 50), f"Customer Visit Itinerary {_campaign(index)}", fontsize=20)
    slots = [
        ("08:15", "Arrival briefing", f"Discuss {_theme(index)} and onboarding friction."),
        ("10:30", "Product workshop", "Validate one-click page save and screenshot capture expectations."),
        ("13:00", "Working lunch", "Review token reduction evidence and benchmark narrative."),
        ("15:45", "Ops review", "Confirm parser fallback, OCR behavior, and table fidelity requirements."),
    ]
    top = 108
    for offset, (time_label, title, note) in enumerate(slots):
        y = top + (offset * 120)
        page.draw_rect((48, y, 542, y + 92), width=1)
        page.insert_text((62, y + 24), time_label, fontsize=16)
        page.insert_text((154, y + 24), title, fontsize=14)
        page.insert_textbox((154, y + 38, 520, y + 82), textwrap.fill(note, width=48), fontsize=11, lineheight=1.25)
    _save_pdf(doc, path)
    return path


def _pdf_knowledge_export(path: Path, index: int) -> Path:
    doc = _new_pdf()
    page = doc.new_page(width=595, height=842)
    page.insert_text((48, 52), f"Knowledge Base Export {_campaign(index)}", fontsize=20)
    sections = [
        ("Checklist", "- verify indexed chunks exist\n- verify queries return hits\n- compare first-turn and warm-cache token counts"),
        ("Troubleshooting", f"If OpenDataLoader succeeds for {_team(index)} but emits logs, parse the JSON payload after the JVM preamble."),
        ("Policy", "Do not compare token-reduction headline numbers unless the retrieval methodology and quality bar match."),
    ]
    y = 126
    for title, body in sections:
        page.draw_rect((48, y, 542, y + 132), width=1)
        page.insert_text((60, y + 22), title, fontsize=14)
        page.insert_textbox((60, y + 38, 528, y + 118), body, fontsize=11, lineheight=1.35)
        y += 150
    _save_pdf(doc, path)
    return path


def _pdf_meeting_agenda(path: Path, index: int) -> Path:
    doc = _new_pdf()
    page = doc.new_page(width=612, height=792)
    page.insert_text((54, 54), f"Launch Review Agenda {_campaign(index)}", fontsize=22)
    bullets = [
        f"Benchmark theme: {_theme(index)}",
        "Replace toy repo-file timing with user-shaped corpus evidence",
        "Keep browser capture simple and desktop-provisioned",
        "Separate first-turn token savings from warm-cache savings",
        "Track parser quality on tables, scans, and exports",
    ]
    y = 112
    for bullet in bullets:
        page.insert_text((70, y), f"- {bullet}", fontsize=13)
        y += 36
    note = textwrap.fill(
        f"The {_team(index)} team in {_region(index)} wants evidence that repeated follow-up questions become cheaper without losing citations.",
        width=66,
    )
    page.insert_textbox((54, 330, 548, 480), note, fontsize=12, lineheight=1.35)
    _save_pdf(doc, path)
    return path


def _pdf_risk_register(path: Path, index: int) -> Path:
    doc = _new_pdf()
    page = doc.new_page(width=612, height=792)
    page.insert_text((48, 48), f"Risk Register {_campaign(index)}", fontsize=21)
    headers = ["Risk", "Impact", "Mitigation"]
    top = 110
    cols = [48, 220, 340, 560]
    heights = [110, 210, 310, 410, 510]
    for x in cols:
        page.draw_line((x, top), (x, 580), width=1)
    for y in [top, *heights, 580]:
        page.draw_line((48, y), (560, y), width=1)
    for idx, header in enumerate(headers):
        page.insert_text((cols[idx] + 8, top + 20), header, fontsize=11)
    rows = [
        (f"{_theme(index)} benchmark is empty", "Misleading claims", "Reject zero-hit and zero-chunk runs"),
        ("OCR backlog causes retry storm", "Delayed ingestion", "Throttle retries and isolate worker"),
        ("Invoice tables flatten", "Broken cost retrieval", "Track PDF layout fidelity and parser choice"),
        ("Manual extension setup leaks complexity", "Onboarding dropoff", "Provision config from desktop"),
    ]
    for row_index, row in enumerate(rows):
        y = top + 40 + (row_index * 100)
        page.insert_textbox((56, y, 212, y + 72), row[0], fontsize=10)
        page.insert_textbox((228, y, 332, y + 72), row[1], fontsize=10)
        page.insert_textbox((348, y, 552, y + 72), row[2], fontsize=10)
    _save_pdf(doc, path)
    return path


def _pdf_product_memo(path: Path, index: int) -> Path:
    doc = _new_pdf()
    page = doc.new_page(width=612, height=792)
    page.draw_rect((42, 42, 570, 748), width=1)
    page.insert_text((58, 62), f"Product Memo {_campaign(index)}", fontsize=20)
    page.insert_textbox(
        (58, 110, 380, 300),
        textwrap.fill(
            f"The {_theme(index)} initiative for {_team(index)} favors a smaller browser extension surface and stronger retrieval-backed answers. "
            "PDF ingestion must keep invoice and report structure usable. "
            "Repeat-turn savings are a product story only when the first turn stays grounded.",
            width=44,
        ),
        fontsize=12,
        lineheight=1.35,
    )
    page.draw_rect((400, 110, 550, 300), width=1)
    page.insert_text((416, 136), "Metrics", fontsize=14)
    metrics = [
        f"Warm-cache goal: {58 + (index % 18)}%",
        f"Import SLA: {1.2 + ((index % 7) * 0.2):.1f}s",
        "Capture setup: one click",
        f"Region: {_region(index)}",
    ]
    y = 170
    for item in metrics:
        page.insert_text((416, y), item, fontsize=11)
        y += 32
    _save_pdf(doc, path)
    return path


def _pdf_faq_sheet(path: Path, index: int) -> Path:
    doc = _new_pdf()
    page = doc.new_page(width=612, height=792)
    page.insert_text((48, 48), f"Benchmark FAQ {_campaign(index)}", fontsize=21)
    faqs = [
        ("Why user-shaped files?", "Because repo test files do not prove import timing, parser robustness, or grounded retrieval behavior."),
        ("Why separate first-turn and warm-cache savings?", "Because they are different mechanisms with different product claims."),
        ("What does this artifact track?", f"It tracks {_theme(index)}, parser fidelity, and supportable browser capture for {_team(index)}."),
        ("What invalidates a run?", "Zero indexed chunks, zero retrieval hits, or a corpus that does not represent real user file types."),
    ]
    y = 106
    for question, answer in faqs:
        page.insert_text((52, y), question, fontsize=14)
        page.insert_textbox((52, y + 14, 552, y + 78), textwrap.fill(answer, width=78), fontsize=11, lineheight=1.3)
        y += 118
    _save_pdf(doc, path)
    return path


def _create_docx_variant(path: Path, index: int) -> Path:
    from docx import Document

    document = Document()
    document.add_heading(f"Implementation Plan {_campaign(index)}", level=1)
    document.add_paragraph(
        f"The {_team(index)} team uses this document to track {_theme(index)} benchmarks, retrieval quality, and browser capture simplification."
    )
    document.add_paragraph(
        "Phase one proves ingestion timing on mixed file types. Phase two verifies retrieval-backed token reduction. Phase three hardens packaging and operator evidence."
    )
    table = document.add_table(rows=0, cols=3)
    rows = [
        ["Area", "Metric", "Target"],
        ["PDF parser", "success rate", ">= 95 percent"],
        ["Ingestion", "median latency", "< 2.5 seconds"],
        ["Context", "warm-cache reduction", ">= 50 percent"],
        ["Capture", "setup friction", "one click"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    document.save(path)
    return path


def _markdown_body(index: int) -> str:
    return textwrap.dedent(
        f"""
        ---
        owner: {_team(index)}
        region: {_region(index)}
        campaign: {_campaign(index)}
        ---

        # Note {_campaign(index)}

        The current theme is {_theme(index)}.

        Users trust the product when imports stay fast, search answers cite original artifacts, and repeated follow-up questions cost less than the first answer.

        ## Action items

        - keep browser capture one-click and desktop-provisioned
        - preserve PDF table fidelity for invoices and reports
        - reject empty retrieval-hit benchmarks before publishing token claims
        """
    ).strip() + "\n"


def _text_body(index: int) -> str:
    return textwrap.dedent(
        f"""
        Sync log {_campaign(index)}
        team={_team(index)}
        region={_region(index)}
        theme={_theme(index)}
        summary=Repeated screenshot retries can flood the OCR backlog if recovery throttling is missing.
        policy=Warm-cache token savings should be reported separately from first-turn packet reduction.
        """
    ).strip() + "\n"


def _html_title(index: int) -> str:
    return f"Dashboard Export {_campaign(index)}"


def _html_body(index: int) -> str:
    items = [
        f"<li>Theme: {_theme(index)}</li>",
        f"<li>Team: {_team(index)}</li>",
        "<li>Repeat-turn cache reduction remains a headline metric.</li>",
        "<li>PDF retrieval quality depends on preserving tables and reading order.</li>",
        "<li>Browser capture should stay simple enough for non-technical operators.</li>",
    ]
    return (
        f"<h1>{_html_title(index)}</h1>"
        f"<p>Campaign {_campaign(index)} tracks retrieval trust, parser timing, and supportable capture workflows in {_region(index)}.</p>"
        f"<ul>{''.join(items)}</ul>"
    )


def _json_payload(index: int) -> dict:
    return {
        "campaign": _campaign(index),
        "team": _team(index),
        "region": _region(index),
        "theme": _theme(index),
        "observations": [
            "repeat-turn token savings matter",
            "browser capture setup should stay simple",
            "parser benchmarks need retrieval-backed evidence",
        ],
        "targets": {
            "first_turn_reduction_percent": 35 + (index % 12),
            "warm_cache_reduction_percent": 58 + (index % 25),
        },
    }


def _csv_rows(index: int) -> list[list[str]]:
    return [
        ["segment", "first_turn_tokens", "repeat_turn_tokens", "imports_per_week", "theme"],
        [_team(index), str(5200 + (index % 900)), str(1700 + (index % 500)), str(20 + (index % 40)), _theme(index)],
        [_team(index + 1), str(5600 + (index % 850)), str(1800 + (index % 550)), str(24 + (index % 30)), _theme(index + 1)],
        [_team(index + 2), str(6100 + (index % 800)), str(2100 + (index % 450)), str(30 + (index % 35)), _theme(index + 2)],
    ]


def _write_markdown(path: Path, text: str) -> Path:
    path.write_text(text, encoding="utf-8")
    return path


def _write_text(path: Path, text: str) -> Path:
    path.write_text(text, encoding="utf-8")
    return path


def _write_html(path: Path, title: str, body: str) -> Path:
    payload = f"""<!doctype html>
<html lang="en">
<head><meta charset="utf-8" /><title>{title}</title></head>
<body>{body}</body>
</html>
"""
    path.write_text(payload, encoding="utf-8")
    return path


def _write_json(path: Path, payload: dict) -> Path:
    path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    return path


def _write_csv(path: Path, *, rows: list[list[str]]) -> Path:
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerows(rows)
    return path


def _new_pdf():
    import fitz

    return fitz.open()


def _save_pdf(doc, path: Path) -> None:
    doc.save(path)
    doc.close()
