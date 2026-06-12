from pathlib import Path
import os
import re
from html.parser import HTMLParser
from urllib.parse import urljoin, urlparse
from urllib.request import HTTPRedirectHandler, Request, build_opener

from backend.app.core.network_security import (
    NetworkSecurityError,
    strip_url_credentials,
    validate_public_http_url,
    validate_response_peer,
)
from backend.app.core.ocr import OCRError, ocr_image, ocr_pdf_pages


SUPPORTED_TEXT_EXTENSIONS = {
    ".asc",
    ".csv",
    ".htm",
    ".html",
    ".json",
    ".jsonl",
    ".log",
    ".md",
    ".markdown",
    ".rtf",
    ".text",
    ".tsv",
    ".txt",
    ".xml",
    ".yaml",
    ".yml",
}
SUPPORTED_CODE_EXTENSIONS = {
    ".bat",
    ".c",
    ".cpp",
    ".cs",
    ".css",
    ".go",
    ".java",
    ".js",
    ".jsx",
    ".kt",
    ".lua",
    ".php",
    ".ps1",
    ".py",
    ".rb",
    ".rs",
    ".sh",
    ".sql",
    ".swift",
    ".toml",
    ".ts",
    ".tsx",
}
SUPPORTED_IMAGE_EXTENSIONS = {".bmp", ".gif", ".jpeg", ".jpg", ".png", ".tif", ".tiff", ".webp"}
SUPPORTED_MEDIA_EXTENSIONS = {".aac", ".flac", ".m4a", ".mov", ".mp3", ".mp4", ".ogg", ".wav", ".webm"}
SUPPORTED_DOCUMENT_EXTENSIONS = (
    SUPPORTED_TEXT_EXTENSIONS
    | SUPPORTED_CODE_EXTENSIONS
    | SUPPORTED_IMAGE_EXTENSIONS
    | SUPPORTED_MEDIA_EXTENSIONS
    | {".docx", ".pdf"}
)
MAX_LOCAL_FILE_BYTES = 50 * 1024 * 1024
MAX_LOCAL_MEDIA_BYTES = 250 * 1024 * 1024
MAX_LINK_BYTES = 2_000_000
MAX_REDIRECTS = 5


class ExtractionError(Exception):
    pass


def extract_text_from_path(path: str) -> tuple[str, str]:
    title, pages = extract_pages_from_path(path)
    return title, "\n\n".join(page for page in pages if page.strip()).strip()


def extract_pages_from_path(path: str) -> tuple[str, list[str]]:
    if not os.environ.get("CML_PARSER_WORKER"):
        from backend.app.core.quarantine import parse_candidate_file, validate_candidate_file

        candidate = validate_candidate_file(path)
        parsed = parse_candidate_file(candidate)
        return parsed["title"], parsed["pages"]
    return extract_pages_from_validated_path(path)


def extract_pages_from_validated_path(path: str) -> tuple[str, list[str]]:
    source_path = Path(path).expanduser()
    if not source_path.exists() or not source_path.is_file():
        raise ExtractionError("File does not exist or is not readable")
    suffix = source_path.suffix.lower()
    try:
        size = source_path.stat().st_size
        max_bytes = MAX_LOCAL_MEDIA_BYTES if suffix in SUPPORTED_MEDIA_EXTENSIONS else MAX_LOCAL_FILE_BYTES
        if size > max_bytes:
            raise ExtractionError("File is too large to ingest safely")
    except OSError as exc:
        raise ExtractionError("File does not exist or is not readable") from exc

    if suffix in SUPPORTED_TEXT_EXTENSIONS:
        if suffix in {".md", ".markdown"}:
            return source_path.name, [_extract_markdown_text(source_path)]
        return source_path.name, [_extract_plain_text(source_path)]
    if suffix in SUPPORTED_CODE_EXTENSIONS:
        return source_path.name, [_extract_code_text(source_path)]
    if suffix == ".docx":
        return source_path.name, [_extract_docx_text(source_path)]
    if suffix == ".pdf":
        return source_path.name, _extract_pdf_pages(source_path)
    if suffix in SUPPORTED_IMAGE_EXTENSIONS:
        return source_path.name, [_extract_image_text(source_path)]
    if suffix in SUPPORTED_MEDIA_EXTENSIONS:
        return source_path.name, [_extract_media_metadata(source_path)]

    raise ExtractionError("This file type is not supported for local vault ingestion yet")


def _extract_plain_text(source_path: Path) -> str:
    try:
        text = source_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            text = source_path.read_text(encoding="utf-8-sig")
        except UnicodeDecodeError:
            text = source_path.read_text(encoding="cp1252")
    except OSError as exc:
        raise ExtractionError(str(exc)) from exc
    return _clean_text_payload(source_path, text)


def _extract_code_text(source_path: Path) -> str:
    text = _extract_plain_text(source_path)
    return f"Code file: {source_path.name}\n\n{text}"


def _extract_markdown_text(source_path: Path) -> str:
    text = _extract_plain_text(source_path)
    body, frontmatter = _split_frontmatter(text)
    metadata: list[str] = []
    if frontmatter:
        metadata.append("Obsidian/frontmatter metadata:")
        metadata.extend(f"- {line.strip()}" for line in frontmatter.splitlines() if line.strip())

    wiki_links = sorted(set(re.findall(r"(?<!!)\[\[([^\]\n]+)\]\]", body)))
    embeds = sorted(set(re.findall(r"!\[\[([^\]\n]+)\]\]", body)))
    markdown_links = sorted(set(match[1] for match in re.findall(r"\[([^\]\n]+)\]\(([^)\n]+)\)", body)))

    if wiki_links:
        metadata.append("Wiki links: " + ", ".join(wiki_links[:50]))
    if embeds:
        metadata.append("Embedded attachments: " + ", ".join(embeds[:50]))
    if markdown_links:
        metadata.append("Markdown links: " + ", ".join(markdown_links[:50]))

    if not metadata:
        return text
    return f"{body.strip()}\n\n" + "\n".join(metadata)


def _split_frontmatter(text: str) -> tuple[str, str]:
    if not text.startswith("---\n"):
        return text, ""
    end = text.find("\n---", 4)
    if end == -1:
        return text, ""
    frontmatter = text[4:end].strip()
    body = text[end + len("\n---"):].strip()
    return body, frontmatter


def _extract_docx_text(source_path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ExtractionError("DOCX extraction requires python-docx to be installed") from exc

    try:
        document = Document(str(source_path))
    except Exception as exc:
        raise ExtractionError(f"Could not read DOCX file: {exc}") from exc

    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    text = "\n".join(paragraphs).strip()
    if not text:
        raise ExtractionError("No readable text was found in this DOCX file")
    return text


def _extract_pdf_text(source_path: Path) -> str:
    pages = _extract_pdf_pages(source_path)
    return "\n\n".join(page for page in pages if page.strip()).strip()


def _extract_pdf_pages(source_path: Path) -> list[str]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ExtractionError("PDF extraction requires pypdf to be installed") from exc

    try:
        reader = PdfReader(str(source_path))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except Exception as exc:
        raise ExtractionError(f"Could not read PDF file: {exc}") from exc

    readable_pages = [page for page in pages if page.strip()]
    if not readable_pages:
        try:
            ocr_pages = ocr_pdf_pages(source_path)
        except OCRError as exc:
            raise ExtractionError(f"No readable text was found in this PDF file. {exc}") from exc
        readable_ocr_pages = [page for page in ocr_pages if page.strip()]
        if not readable_ocr_pages:
            raise ExtractionError("No readable text was found in this PDF file, including after local OCR.")
        return ocr_pages
    return pages


def _extract_image_text(source_path: Path) -> str:
    try:
        text = ocr_image(source_path)
        if text.strip():
            return text
    except OCRError:
        pass
    return _file_metadata_text(source_path, note="Image stored in vault metadata. OCR runtime is not configured yet.")


def _extract_media_metadata(source_path: Path) -> str:
    return _file_metadata_text(
        source_path,
        note="Media file stored in vault metadata. Audio/video transcription is not configured yet.",
    )


def _file_metadata_text(source_path: Path, *, note: str) -> str:
    stat = source_path.stat()
    return "\n".join(
        [
            note,
            f"File name: {source_path.name}",
            f"File type: {source_path.suffix.lower() or 'unknown'}",
            f"Size bytes: {stat.st_size}",
        ]
    )


def _clean_text_payload(source_path: Path, text: str) -> str:
    cleaned = text.strip()
    if not cleaned:
        raise ExtractionError(f"No readable text was found in {source_path.name}")
    return cleaned


class _TextHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []
        self.page_title = ""
        self.meta_title = ""
        self.cover_image_url = ""
        self._skip_depth = 0
        self._in_title = False

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attrs_by_name = {name.lower(): value or "" for name, value in attrs}
        if tag in {"script", "style", "noscript"}:
            self._skip_depth += 1
        if tag == "title":
            self._in_title = True
        if tag == "meta":
            key = attrs_by_name.get("property") or attrs_by_name.get("name")
            content = attrs_by_name.get("content", "").strip()
            if key in {"og:title", "twitter:title"} and content and not self.meta_title:
                self.meta_title = content
            if key in {"og:image", "twitter:image"} and content and not self.cover_image_url:
                self.cover_image_url = content
        if tag in {"p", "br", "li", "h1", "h2", "h3", "h4", "h5", "h6"}:
            self._parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "noscript"} and self._skip_depth > 0:
            self._skip_depth -= 1
        if tag == "title":
            self._in_title = False
        if tag in {"p", "li", "div", "section", "article"}:
            self._parts.append("\n")

    def handle_data(self, data: str) -> None:
        if self._in_title:
            title = " ".join(data.split())
            if title and not self.page_title:
                self.page_title = title
        if self._skip_depth == 0:
            cleaned = " ".join(data.split())
            if cleaned:
                self._parts.append(cleaned)

    def text(self) -> str:
        return "\n".join(part for part in self._parts if part.strip()).strip()


def extract_text_from_url(url: str) -> tuple[str, str, str | None]:
    title, text, cover_image_url, _security = extract_text_from_url_with_security(url)
    return title, text, cover_image_url


def extract_text_from_url_with_security(url: str) -> tuple[str, str, str | None, dict]:
    from backend.app.core.browser_ingestion import static_web_security

    url = strip_url_credentials(url)
    try:
        validate_public_http_url(url)
    except NetworkSecurityError as exc:
        raise ExtractionError(str(exc)) from exc

    request = Request(
        url,
        headers={
            "User-Agent": "CML/0.1 local context ingestion",
            "Accept": "text/html,text/plain;q=0.9,*/*;q=0.8",
        },
    )

    try:
        response, final_url = _safe_open(request, timeout=12)
        with response:
            content_type = response.headers.get("content-type", "")
            body = response.read(MAX_LINK_BYTES + 1)
    except OSError as exc:
        raise ExtractionError(str(exc)) from exc
    if len(body) > MAX_LINK_BYTES:
        raise ExtractionError("Link response is too large to ingest safely")

    parsed = urlparse(final_url)
    try:
        decoded = body.decode("utf-8")
    except UnicodeDecodeError:
        decoded = body.decode("utf-8", errors="replace")

    security = static_web_security(final_url)
    if "html" in content_type:
        parser = _TextHTMLParser()
        parser.feed(decoded)
        text = parser.text()
        title = parser.meta_title or parser.page_title
        cover_image_url = urljoin(url, parser.cover_image_url) if parser.cover_image_url else None
        if _needs_dynamic_extraction(text, decoded):
            dynamic = _extract_dynamic_text_from_url(final_url)
            if dynamic is not None and len(dynamic[1]) > len(text):
                title, text, dynamic_cover = dynamic[:3]
                cover_image_url = dynamic_cover or cover_image_url
                if len(dynamic) > 3:
                    security = dynamic[3]
    else:
        text = decoded.strip()
        title = ""
        cover_image_url = None

    if not text:
        raise ExtractionError("No readable text was found at this link")

    fallback_title = (parsed.netloc + parsed.path).rstrip("/") or url
    return title or fallback_title, text, cover_image_url, security


def link_extraction_diagnostics(url: str) -> dict:
    sanitized_url = strip_url_credentials(url)
    diagnostics = {
        "input_url_had_credentials": sanitized_url != url,
        "sanitized_url": sanitized_url,
        "allowed": False,
        "security_error": "",
        "static_timeout_seconds": 12,
        "dynamic_timeout_seconds": 12,
        "max_response_bytes": MAX_LINK_BYTES,
        "dynamic_fallback_available": False,
        "browser_isolation": {},
        "extraction_order": ["static_http", "isolated_browser_rendered_dynamic_fallback"],
        "quality_classification": "unknown",
        "quality_reason": "Fetch was not run during diagnostics.",
    }
    try:
        validate_public_http_url(sanitized_url)
        diagnostics["allowed"] = True
        diagnostics["quality_classification"] = "diagnostic_only"
    except NetworkSecurityError as exc:
        diagnostics["security_error"] = str(exc)
        diagnostics["quality_classification"] = "blocked_by_security"
        diagnostics["quality_reason"] = str(exc)
    from backend.app.core.browser_ingestion import browser_ingestion_diagnostics

    browser_diagnostics = browser_ingestion_diagnostics()
    diagnostics["browser_isolation"] = browser_diagnostics
    diagnostics["dynamic_fallback_available"] = bool(browser_diagnostics["available"])
    if diagnostics["allowed"] and diagnostics["dynamic_fallback_available"]:
        diagnostics["quality_reason"] = "Static extraction is attempted first; browser fallback is available for thin dynamic pages."
    elif diagnostics["allowed"]:
        diagnostics["quality_reason"] = "Static extraction is available; browser fallback runtime is not installed."
    return diagnostics


def _needs_dynamic_extraction(text: str, html: str) -> bool:
    lowered = html.lower()
    script_count = lowered.count("<script")
    app_markers = ("id=\"root\"", "id=\"app\"", "__next", "data-reactroot", "vite")
    return (len(text.strip()) < 500 and script_count >= 3) or any(marker in lowered for marker in app_markers)


def _extract_dynamic_text_from_url(url: str) -> tuple[str, str, str | None, dict] | None:
    try:
        from backend.app.core.browser_ingestion import extract_dynamic_text_from_url_isolated

        return extract_dynamic_text_from_url_isolated(url)
    except Exception:
        return None


class _NoRedirectHandler(HTTPRedirectHandler):
    def redirect_request(self, req, fp, code, msg, headers, newurl):
        return None


def _safe_open(request: Request, timeout: int):
    opener = build_opener(_NoRedirectHandler)
    current = request
    for _ in range(MAX_REDIRECTS + 1):
        try:
            response = opener.open(current, timeout=timeout)
            try:
                validate_public_http_url(response.geturl())
                validate_response_peer(response)
            except NetworkSecurityError as validation_exc:
                raise ExtractionError(str(validation_exc)) from validation_exc
            return response, response.geturl()
        except Exception as exc:
            code = getattr(exc, "code", None)
            headers = getattr(exc, "headers", {})
            if code not in {301, 302, 303, 307, 308}:
                raise
            location = headers.get("Location")
            if not location:
                raise ExtractionError("Link redirect did not include a target") from exc
            next_url = urljoin(current.full_url, location)
            next_url = strip_url_credentials(next_url)
            try:
                validate_public_http_url(next_url)
            except NetworkSecurityError as validation_exc:
                raise ExtractionError(str(validation_exc)) from validation_exc
            current = Request(next_url, headers=dict(current.header_items()))
    raise ExtractionError("Too many redirects while fetching link")
